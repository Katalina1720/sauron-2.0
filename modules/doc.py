from docx import Document
from PIL import Image
from io import BytesIO

from modules.img_plugin import pillow_to_txt


# извлечение текста из картинок
def extract_images(doc: Document) -> str:
    text = []
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            img = rel.target_part.blob
            image = Image.open(BytesIO(img))
            text.append(pillow_to_txt(image))
    return "".join(text)


def extract_tables(doc: Document) -> list[str]:
    text = []
    for table in doc.tables:
        for i, _ in enumerate(table.rows):
            for j, __ in enumerate(table.columns):
                text.append(table.cell(i, j).text)
    return text


def read_file(docx_path) -> str:
    document = Document(docx_path)
    text = []
    for paragraph in document.paragraphs:
        text.append(paragraph.text)
    text += extract_tables(document)
    text += extract_images(document)

    return "\n".join(text)
